"""Извлечение текста из документов для интейка.

Поддержка: txt/md, csv, xlsx/xlsm, docx, pdf. Форматы подключаются
лениво: тяжёлые зависимости импортируются только при обращении, а их
отсутствие превращается в понятную ошибку, не роняя всё приложение.
"""
from __future__ import annotations

import csv
import io
from dataclasses import dataclass, field

from .errors import IntakeError

MAX_CHARS = 200_000  # защита промпта от гигантских документов


@dataclass
class ExtractedDoc:
    """Результат извлечения: имя, тип, текст и метаданные."""

    filename: str
    kind: str                       # txt | csv | xlsx | docx | pdf
    text: str
    meta: dict = field(default_factory=dict)

    @property
    def truncated(self) -> bool:
        return bool(self.meta.get("truncated"))


def _clip(text: str, meta: dict) -> str:
    if len(text) > MAX_CHARS:
        meta["truncated"] = True
        return text[:MAX_CHARS]
    return text


def _decode(raw: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1251"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _extract_txt(raw: bytes, meta: dict) -> str:
    return _decode(raw)


def _extract_csv(raw: bytes, meta: dict) -> str:
    text = _decode(raw)
    try:
        dialect = csv.Sniffer().sniff(text[:2000], delimiters=";,\t")
    except csv.Error:
        dialect = csv.excel
    rows = list(csv.reader(io.StringIO(text), dialect))
    meta["rows"] = len(rows)
    return "\n".join(" | ".join(cell.strip() for cell in row) for row in rows if any(row))


def _extract_xlsx(raw: bytes, meta: dict) -> str:
    try:
        import openpyxl
    except ImportError as exc:  # pragma: no cover
        raise IntakeError("для чтения xlsx требуется пакет openpyxl") from exc
    workbook = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    parts: list[str] = []
    meta["sheets"] = workbook.sheetnames
    for sheet in workbook.worksheets:
        parts.append(f"=== Лист: {sheet.title} ===")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(v) for v in row if v is not None and str(v).strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_docx(raw: bytes, meta: dict) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise IntakeError("для чтения docx требуется пакет python-docx") from exc
    document = docx.Document(io.BytesIO(raw))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    meta["paragraphs"] = len(parts)
    return "\n".join(parts)


def _extract_pdf(raw: bytes, meta: dict) -> str:
    try:
        import pdfplumber
    except ImportError as exc:  # pragma: no cover
        raise IntakeError("для чтения pdf требуется пакет pdfplumber") from exc
    parts: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            meta["pages"] = len(pdf.pages)
            for number, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    parts.append(f"=== Страница {number} ===\n{text}")
    except Exception as exc:
        raise IntakeError(f"не удалось разобрать PDF: {exc}") from exc
    if not parts:
        raise IntakeError(
            "в PDF не найден текстовый слой (вероятно, скан); потребуется OCR — этот режим пока не поддержан"
        )
    return "\n".join(parts)


_DISPATCH = {
    "txt": _extract_txt, "md": _extract_txt,
    "csv": _extract_csv,
    "xlsx": _extract_xlsx, "xlsm": _extract_xlsx,
    "docx": _extract_docx,
    "pdf": _extract_pdf,
}


def extract_text(filename: str, raw: bytes) -> ExtractedDoc:
    """Извлекает текст из документа по расширению имени файла."""
    if not raw:
        raise IntakeError(f"файл {filename!r} пуст")
    ext = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()
    handler = _DISPATCH.get(ext)
    if handler is None:
        supported = ", ".join(sorted(set(_DISPATCH)))
        raise IntakeError(f"формат .{ext or '?'} не поддерживается (доступны: {supported})")
    meta: dict = {}
    text = _clip(handler(raw, meta), meta)
    if not text.strip():
        raise IntakeError(f"из файла {filename!r} не удалось извлечь текст")
    return ExtractedDoc(filename=filename, kind="txt" if ext == "md" else ext, text=text, meta=meta)
