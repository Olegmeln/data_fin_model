"""Тесты интейка: textract по форматам и экстрактор с инъектированным LLM."""
import io
import json

import pytest

from app.finmodel.intake import ExtractedDoc, IntakeError, extract_assumptions, extract_text
from app.finmodel.intake.extractor import build_prompt
from conftest import SAMPLES_DIR


# --------------------------------------------------------------- textract

class TestTextract:
    def test_txt_fixture(self):
        raw = (SAMPLES_DIR / "бизнес_план_мини_demo.txt").read_bytes()
        doc = extract_text("бизнес_план_мини_demo.txt", raw)
        assert doc.kind == "txt"
        assert "пастила" in doc.text.lower()
        assert not doc.truncated

    def test_csv(self):
        raw = "Статья;Сумма\nОборудование;65\nРемонт;40\n".encode("cp1251")
        doc = extract_text("смета.csv", raw)
        assert doc.kind == "csv"
        assert "Оборудование | 65" in doc.text
        assert doc.meta["rows"] == 3

    def test_xlsx(self):
        import openpyxl
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "Смета"
        sheet.append(["Статья", "Сумма, млн"])
        sheet.append(["Оборудование", 65])
        buffer = io.BytesIO()
        workbook.save(buffer)
        doc = extract_text("смета.xlsx", buffer.getvalue())
        assert doc.kind == "xlsx"
        assert doc.meta["sheets"] == ["Смета"]
        assert "Оборудование | 65" in doc.text

    def test_docx(self):
        import docx
        document = docx.Document()
        document.add_paragraph("Инвестиции: 120 млн руб.")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Кредит"
        table.rows[0].cells[1].text = "90 млн"
        buffer = io.BytesIO()
        document.save(buffer)
        doc = extract_text("план.docx", buffer.getvalue())
        assert doc.kind == "docx"
        assert "Инвестиции: 120 млн руб." in doc.text
        assert "Кредит | 90 млн" in doc.text

    def test_unknown_extension(self):
        with pytest.raises(IntakeError, match="не поддерживается"):
            extract_text("файл.xyz", b"data")

    def test_empty_file(self):
        with pytest.raises(IntakeError, match="пуст"):
            extract_text("план.txt", b"")

    def test_broken_pdf(self):
        with pytest.raises(IntakeError):
            extract_text("скан.pdf", b"%PDF-1.4 broken broken")

    def test_truncation(self):
        raw = ("строка\n" * 60_000).encode()
        doc = extract_text("большой.txt", raw)
        assert doc.truncated
        assert len(doc.text) <= 200_000


# --------------------------------------------------------------- extractor

def _valid_payload() -> dict:
    return {
        "schema": "finmodel.assumptions.v1",
        "profile": {"name": "Пастила «Яблоневый сад»", "industry": "пищевое производство",
                    "horizon_years": 7, "model_start": "2026-01-01"},
        "products": [
            {"name": "Пастила классическая", "unit": "кг", "start_price": 900},
            {"name": "Подарочные наборы", "kind": "goods", "unit": "шт", "start_price": 1500},
        ],
        "capex": {"items": [{"name": "Ремонт цеха", "amount": 40}, {"name": "Оборудование", "amount": 65}],
                  "depreciation_months": 84},
        "financing": {"equity_amount": 30, "facilities": [
            {"name": "Инвестиционный", "amount": 90, "term_months": 60, "grace_months": 12,
             "rate": {"points": [{"value_pct": 12}]}}]},
        "taxes": {"regime": "УСН 15%", "vat": {"points": [{"value_pct": 0}]},
                  "profit": {"points": [{"value_pct": 15}]}},
        "valuation": {"discount_rate_pct": 14},
        "open_questions": [{"question": "Аренда цеха или выкуп?", "severity": "blocker"}],
        "sources": {"financing.facilities.0.amount": {
            "method": "extracted", "document": "бизнес_план_мини_demo.txt",
            "locator": "раздел Финансирование", "confidence": 0.95}},
    }


def _docs() -> list[ExtractedDoc]:
    raw = (SAMPLES_DIR / "бизнес_план_мини_demo.txt").read_bytes()
    return [extract_text("бизнес_план_мини_demo.txt", raw)]


class TestExtractor:
    def test_happy_path_with_fake_llm(self):
        calls = {}

        def fake_llm(system: str, user: str) -> str:
            calls["system"], calls["user"] = system, user
            return json.dumps(_valid_payload(), ensure_ascii=False)

        aset = extract_assumptions(_docs(), llm=fake_llm)
        assert aset.profile.horizon_years == 7
        assert {p.name for p in aset.products} == {"Пастила классическая", "Подарочные наборы"}
        assert aset.financing.facilities[0].grace_months == 12
        assert aset.open_questions[0].severity == "blocker"
        # промпт содержит схему и текст документа
        assert "finmodel.assumptions.v1" in calls["system"]
        assert "пастила" in calls["user"].lower()

    def test_markdown_fences_stripped(self):
        def fenced_llm(system: str, user: str) -> str:
            return "```json\n" + json.dumps(_valid_payload(), ensure_ascii=False) + "\n```"

        aset = extract_assumptions(_docs(), llm=fenced_llm)
        assert aset.profile.name.startswith("Пастила")

    def test_non_json_answer(self):
        with pytest.raises(IntakeError, match="не-JSON"):
            extract_assumptions(_docs(), llm=lambda s, u: "к сожалению, не могу")

    def test_schema_violation_reported(self):
        def bad_llm(system: str, user: str) -> str:
            payload = _valid_payload()
            payload["schema"] = "finmodel.assumptions.v99"
            return json.dumps(payload)

        with pytest.raises(IntakeError, match="не прошёл схему"):
            extract_assumptions(_docs(), llm=bad_llm)

    def test_no_documents(self):
        with pytest.raises(IntakeError, match="ни одного документа"):
            extract_assumptions([], llm=lambda s, u: "{}")

    def test_no_api_key_clear_error(self, monkeypatch):
        from app.config import settings
        monkeypatch.setattr(settings, "ANTHROPIC_API_KEY", "")
        with pytest.raises(IntakeError, match="ANTHROPIC_API_KEY"):
            extract_assumptions(_docs())  # без инъекции — боевой клиент

    def test_prompt_marks_truncated_docs(self):
        doc = ExtractedDoc(filename="big.txt", kind="txt", text="x", meta={"truncated": True})
        _, user = build_prompt([doc])
        assert "обрезан по лимиту" in user
